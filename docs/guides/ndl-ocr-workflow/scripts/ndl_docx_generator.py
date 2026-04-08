#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
NDL DOCX Generator (Canonical Version)
=====================================
国立国会図書館（NDL）から取得したテキスト資料から、
高品質なDOCX（18pt本文 / 22pt見出し / MS明朝・ゴシック）を生成します。

使い方:
    python ndl_docx_generator.py --input _full_ocr_text.txt --output result.docx --title "タイトル" --author "著者名" --pid "123456"

config.py が同じディレクトリに存在する場合、フォント設定や補正辞書を
config.py から自動的に読み込みます。
"""

import argparse
import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# config.py の読み込み（存在する場合）
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    import config
    DEFAULT_BODY_SIZE    = getattr(config, 'DOCX_BODY_SIZE', 18)
    DEFAULT_HEADING_SIZE = getattr(config, 'DOCX_HEADING_SIZE', 22)
    DEFAULT_LINE_SPACING = getattr(config, 'DOCX_LINE_SPACING', 1.5)
    BODY_FONT            = getattr(config, 'DOCX_BODY_FONT', 'MS Mincho')
    HEADING_FONT         = getattr(config, 'DOCX_HEADING_FONT', 'MS Gothic')
    HEADING_COLOR        = getattr(config, 'DOCX_HEADING_COLOR', (0x00, 0x00, 0x80))
    AI_REPAIR_MAP        = getattr(config, 'AI_REPAIR_MAP', {})
except ImportError:
    DEFAULT_BODY_SIZE = 18
    DEFAULT_HEADING_SIZE = 22
    DEFAULT_LINE_SPACING = 1.5
    BODY_FONT = 'MS Mincho'
    HEADING_FONT = 'MS Gothic'
    HEADING_COLOR = (0x00, 0x00, 0x80)
    AI_REPAIR_MAP = {}


def ai_repair_text(text):
    """テキストの正規化と簡易的な補正"""
    for old, new in AI_REPAIR_MAP.items():
        text = text.replace(old, new)
    text = text.replace("\u3002 ", "\u3002")
    return text

def process_paragraph_merge(text):
    """行を論理的な段落に統合する（見出し判定を含む）"""
    lines = text.split('\n')
    merged_blocks = []
    current_block = ""

    heading_pattern = r'^(第[一二三四五六七八九十壹貳參肆伍陸柒捌玖拾]{1,3}[章編回節])|^(自序|目次|要目索引|解題|序文|後記|付録)'

    for line in lines:
        line = line.strip()
        if not line:
            if current_block:
                merged_blocks.append(("P", current_block))
                current_block = ""
            continue

        if re.match(heading_pattern, line):
            if current_block:
                merged_blocks.append(("P", current_block))
                current_block = ""
            merged_blocks.append(("H", line))
        else:
            current_block += line

    if current_block:
        merged_blocks.append(("P", current_block))

    return merged_blocks

def create_docx(args):
    """引数に基づいてDOCXを生成"""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(args.body_size)
    style.paragraph_format.line_spacing = args.line_spacing
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rFonts = style.element.xpath('.//w:rFonts')[0]
    rFonts.set(qn('w:eastAsia'), BODY_FONT)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    content = input_path.read_text(encoding="utf-8", errors="ignore")

    page_token_pattern = r'---\s*(?:page|p\.)\s*(\d+)\s*---'
    tokens = re.split(page_token_pattern, content)

    doc.add_heading(args.title, 0)
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info = []
    if args.author:
        cover_info.append(f"\u8457\u8005\uff1a{args.author}")
    cover_info.append(f"\n[AI OCR Processed]")
    cover_info.append(f"Body {args.body_size}pt / Heading {args.heading_size}pt")
    if args.pid:
        cover_info.append(f"\nPID: {args.pid}")

    run = p_cover.add_run("\n".join(cover_info))
    run.font.size = Pt(14)
    doc.add_page_break()

    for i in range(1, len(tokens), 2):
        page_id = int(tokens[i])
        page_text = tokens[i + 1].strip()

        if args.limit and page_id > args.limit:
            continue
        if not page_text:
            continue

        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_ref = p_ref.add_run(f"[NDL p.{page_id}]")
        run_ref.font.size = Pt(10)
        run_ref.font.italic = True
        run_ref.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        page_text = ai_repair_text(page_text)
        blocks = process_paragraph_merge(page_text)

        for b_type, b_content in blocks:
            if b_type == "H":
                doc.add_page_break()
                h = doc.add_heading(b_content, level=1)
                for h_run in h.runs:
                    h_run.font.color.rgb = RGBColor(*HEADING_COLOR)
                    h_run.font.size = Pt(args.heading_size)
                    h_run.font.name = HEADING_FONT
                    rFonts_h = h_run._element.xpath('.//w:rFonts')[0]
                    rFonts_h.set(qn('w:eastAsia'), HEADING_FONT)
            else:
                p = doc.add_paragraph(b_content)
                p.paragraph_format.first_line_indent = Pt(args.body_size)

    output_path = Path(args.output)
    doc.save(output_path)
    print(f"Success: Generated DOCX at {output_path}")

def main():
    parser = argparse.ArgumentParser(description="NDL DOCX Generator (Canonical)")
    parser.add_argument("--input", required=True, help="Input text file path")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument("--title", default="NDL Document", help="Document title")
    parser.add_argument("--author", default="", help="Author name")
    parser.add_argument("--pid", default="", help="NDL PID")
    parser.add_argument("--body-size", type=int, default=DEFAULT_BODY_SIZE)
    parser.add_argument("--heading-size", type=int, default=DEFAULT_HEADING_SIZE)
    parser.add_argument("--line-spacing", type=float, default=DEFAULT_LINE_SPACING)
    parser.add_argument("--limit", type=int, help="Stop at this page number")

    args = parser.parse_args()
    create_docx(args)

if __name__ == "__main__":
    main()
